import requests
from bs4 import BeautifulSoup
import time
import pyautogui
from docx import Document

# 사용자입력
keyword = pyautogui.prompt('검색어를 입력하세요')
lastpage = int(pyautogui.prompt('몇 페이지까지 크롤링 할까요?'))

# 워드 문서 생성하기
document = Document()

page_num = 1
for i in range(1, lastpage * 10, 10):
    print(f"{page_num}페이지 크롤링 중입니다.===========================")
    response = requests.get(f'https://search.naver.com/search.naver?where=news&sm=tab_jum&query={keyword}&start={i}')
    # print(response)
    html = response.text
    # print(html)
    soup = BeautifulSoup(html, 'html.parser')
    # print(soup)
    articles = soup.select('div.info_group') # 뉴스 기사 div 10개 추출
    # print(articles) # 리스트 형태로 출력

    for article in articles:
        links = article.select('a.info') # 리스트
        if len(links) >= 2: # 링크가 2개 이상이면
            url = links[1].attrs['href'] # 두번째 링크의 href를 추출
            response = requests.get(url)
            html = response.text
            soup = BeautifulSoup(html, 'html.parser')
            # print(url)
            # 만약 스포츠 뉴스 라면
            if 'sports' in response.url:
                title = soup.select_one('h4.title')
                content = soup.select_one('#newsEndContents')
                # 본문 내용안에 불필요한 div, p 삭제
                divs = content.select('div')
                for div in divs:
                    div.decompose()
                paragraphs = content.select('p')
                for p in paragraphs:
                    p.decompose()
            # 만약 연예 뉴스 라면
            elif 'entertain' in response.url:
                title = soup.select_one('.end_tit')
                content = soup.select_one('#articeBody')
            else:
                title = soup.select_one('#title_area')
                content = soup.select_one('#dic_area')
        
            print("==========링크==========\n", url)
            print("==========제목==========\n", title.text.strip())
            print("==========내용==========\n", content.text.strip())


            # 워드에 제목, 링크, 본문 저장하기
            document.add_heading(title.text.strip(), level=0)
            document.add_paragraph(url)
            document.add_paragraph(content.text.strip())

            time.sleep(0.3)
    page_num = page_num + 1

# 워드 문서 저장하기
document.save(f"C://Users//tig06//Desktop//study//study//inflearn//이것이 진짜 크롤링이다 - 실전편//{keyword}_result.docx")